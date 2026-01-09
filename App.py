from docx import Document
from docx.shared import Pt, Mm

def analyser_msas_template(file_path):
    doc = Document(file_path)
    rapport = []
    
    # 1. Vérification des Marges (Cible: 25mm) 
    section = doc.sections[0]
    marge_haut = section.top_margin.mm
    if round(marge_haut) != 25:
        rapport.append(f"❌ Marges : La marge supérieure est de {marge_haut:.1f}mm au lieu de 25mm. ")
    else:
        rapport.append("✅ Marges : Conformité de 25mm respectée. ")

    # 2. Vérification du Titre (Cible: Times New Roman, 20pt, Gras) [cite: 36, 37, 20]
    premier_paragraphe = doc.paragraphs[0]
    titre_texte = premier_paragraphe.text
    
    # Accès au style du premier "run" (segment de texte)
    if premier_paragraphe.runs:
        font = premier_paragraphe.runs[0].font
        taille = font.size.pt if font.size else "Inconnue"
        is_bold = font.bold
        
        if taille != 20:
            rapport.append(f"❌ Titre : La taille est de {taille}pt au lieu de 20pt. ")
        if not is_bold:
            rapport.append("❌ Titre : Le titre doit être en gras. ")
        if font.name and "Times New Roman" not in font.name:
            rapport.append(f"❌ Police : Police détectée '{font.name}' au lieu de Times New Roman. ")

    # 3. Vérification du Corps de Texte (Cible: 10pt) 
    # On teste un paragraphe au hasard dans l'introduction pour le MVP
    for para in doc.paragraphs:
        if "INTRODUCTION" in para.text.upper():
            intro_idx = doc.paragraphs.index(para)
            if intro_idx + 1 < len(doc.paragraphs):
                test_para = doc.paragraphs[intro_idx + 1]
                if test_para.runs:
                    f_size = test_para.runs[0].font.size
                    if f_size and f_size.pt != 10:
                        rapport.append(f"❌ Corps de texte : Taille de {f_size.pt}pt détectée au lieu de 10pt. ")
            break

    # 4. Vérification du Format Colonne (Cible: 2 colonnes) [cite: 21]
    # Note: python-docx lit difficilement le nombre de colonnes complexes, 
    # mais on peut vérifier si une section spécifique est définie.
    if len(doc.sections) < 2:
        rapport.append("⚠️ Mise en page : Assurez-vous d'utiliser deux colonnes après le résumé. [cite: 21]")

    return rapport

# Utilisation
# resultats = analyser_msas_template("votre_article.docx")
# for r in resultats: print(r)